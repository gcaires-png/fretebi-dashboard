#!/usr/bin/env python3
"""Convert the Markdown sources of docs/ex-tarifario into .docx (python-docx) and .pdf (Chromium).
Usage: python3 _build/build_docs.py [--pdf-all]
By default: every .md -> .docx; the broker checklist and the Factiun memo also -> .pdf.
"""
import re, sys, html, subprocess, json, os, pathlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).resolve().parent.parent
PDF_DEFAULT = {"substancia-checklist-protocolo-sei.md", "memorando-factiun-2026-09-02.md"}

# ---------- tiny markdown block parser ----------
def parse(md):
    lines = md.splitlines(); blocks = []; i = 0
    while i < len(lines):
        l = lines[i]
        if not l.strip(): i += 1; continue
        if l.startswith("```"):
            buf = []; i += 1
            while i < len(lines) and not lines[i].startswith("```"): buf.append(lines[i]); i += 1
            blocks.append(("code", "\n".join(buf))); i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)", l)
        if m: blocks.append(("h", len(m.group(1)), m.group(2).strip())); i += 1; continue
        if re.match(r"^-{3,}\s*$", l): blocks.append(("hr",)); i += 1; continue
        if l.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.match(r"^:?-{2,}:?$", c) for c in cells): rows.append(cells)
                i += 1
            blocks.append(("table", rows)); continue
        if l.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"): buf.append(lines[i][1:].strip()); i += 1
            blocks.append(("quote", " ".join(buf))); continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", l)
        if m:
            items = []
            while i < len(lines):
                m2 = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", lines[i])
                if m2:
                    items.append((len(m2.group(1)) // 2, m2.group(2) != "-" and m2.group(2) != "*", m2.group(3))); i += 1
                elif lines[i].startswith("   ") and items:
                    items[-1] = (items[-1][0], items[-1][1], items[-1][2] + " " + lines[i].strip()); i += 1
                else: break
            blocks.append(("list", items)); continue
        buf = []
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|```|\||>|-{3,}\s*$|\s*([-*]|\d+\.)\s)", lines[i]):
            buf.append(lines[i].strip()); i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
def runs(text):
    for part in INLINE.split(text):
        if not part: continue
        if part.startswith("**"): yield part[2:-2], "b"
        elif part.startswith("*"): yield part[1:-1], "i"
        elif part.startswith("`"): yield part[1:-1], "c"
        else: yield part, ""

# ---------- docx ----------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)

def add_runs(par, text, size=None, color=None):
    for t, k in runs(text):
        r = par.add_run(t)
        if k == "b": r.bold = True
        if k == "i": r.italic = True
        if k == "c": r.font.name = "Consolas"; r.font.size = Pt(9)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor.from_string(color)

def to_docx(blocks, out):
    d = Document()
    for s in d.sections: s.left_margin = s.right_margin = Cm(2); s.top_margin = s.bottom_margin = Cm(2)
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    for b in blocks:
        k = b[0]
        if k == "h":
            lvl = min(b[1], 4); p = d.add_heading(level=lvl if lvl > 0 else 1); add_runs(p, b[2])
            for r in p.runs: r.font.color.rgb = RGBColor(0x0E, 0x4F, 0x63)
        elif k == "p":
            p = d.add_paragraph(); add_runs(p, b[1])
        elif k == "quote":
            p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); add_runs(p, b[1], color="4A5A68")
        elif k == "hr":
            p = d.add_paragraph(); pPr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
            bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6"); bt.set(qn("w:space"), "1"); bt.set(qn("w:color"), "D3DAE1"); bdr.append(bt); pPr.append(bdr)
        elif k == "code":
            p = d.add_paragraph(); r = p.add_run(b[1]); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        elif k == "list":
            for lvl, ordered, txt in b[1]:
                p = d.add_paragraph(style="List Number" if ordered else "List Bullet"); p.paragraph_format.left_indent = Cm(0.6 + 0.6 * lvl); add_runs(p, txt)
        elif k == "table":
            rows = b[1]; ncol = max(len(r) for r in rows)
            t = d.add_table(rows=len(rows), cols=ncol); t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    cell = t.cell(ri, ci); cell.text = ""; par = cell.paragraphs[0]
                    add_runs(par, row[ci] if ci < len(row) else "", size=9)
                    if ri == 0:
                        shade(cell, "E9EDF1")
                        for r in par.runs: r.bold = True
            d.add_paragraph()
    d.save(out)

# ---------- html/pdf ----------
CSS = """
@page{size:A4;margin:18mm 16mm}
body{font-family:'IBM Plex Sans','Segoe UI',Calibri,Arial,sans-serif;font-size:10.5pt;line-height:1.45;color:#152029}
h1{font-size:18pt;color:#0E4F63;margin:0 0 6pt;line-height:1.2}
h2{font-size:13pt;color:#0E4F63;margin:16pt 0 6pt;border-bottom:1px solid #D3DAE1;padding-bottom:2pt}
h3{font-size:11pt;margin:12pt 0 4pt}
table{border-collapse:collapse;width:100%;font-size:9pt;margin:6pt 0 10pt;page-break-inside:auto}
th,td{border:1px solid #C9D1D9;padding:4pt 5pt;vertical-align:top;text-align:left}
th{background:#E9EDF1}
tr{page-break-inside:avoid}
blockquote{margin:6pt 0;padding:6pt 10pt;border-left:3px solid #0E4F63;background:#F2F4F6;color:#4A5A68}
code{font-family:Consolas,Menlo,monospace;font-size:9pt;background:#F2F4F6;padding:0 3px}
pre{font-family:Consolas,Menlo,monospace;font-size:8.5pt;background:#F2F4F6;padding:8pt;white-space:pre-wrap}
hr{border:0;border-top:1px solid #D3DAE1;margin:12pt 0}
li{margin:2pt 0}
.foot{position:fixed;bottom:0;font-size:8pt;color:#7B8A97}
"""
def inline_html(text):
    out = ""
    for t, k in runs(text):
        t = html.escape(t)
        out += {"b": f"<b>{t}</b>", "i": f"<i>{t}</i>", "c": f"<code>{t}</code>"}.get(k, t)
    return out

def to_html(blocks, title):
    h = [f"<!doctype html><html><head><meta charset='utf-8'><title>{html.escape(title)}</title><style>{CSS}</style></head><body>"]
    for b in blocks:
        k = b[0]
        if k == "h": h.append(f"<h{min(b[1],4)}>{inline_html(b[2])}</h{min(b[1],4)}>")
        elif k == "p": h.append(f"<p>{inline_html(b[1])}</p>")
        elif k == "quote": h.append(f"<blockquote>{inline_html(b[1])}</blockquote>")
        elif k == "hr": h.append("<hr>")
        elif k == "code": h.append(f"<pre>{html.escape(b[1])}</pre>")
        elif k == "list":
            tag = "ol" if b[1][0][1] else "ul"; h.append(f"<{tag}>" + "".join(f"<li>{inline_html(t)}</li>" for _, _, t in b[1]) + f"</{tag}>")
        elif k == "table":
            rows = b[1]; h.append("<table><thead><tr>" + "".join(f"<th>{inline_html(c)}</th>" for c in rows[0]) + "</tr></thead><tbody>")
            for r in rows[1:]: h.append("<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in r) + "</tr>")
            h.append("</tbody></table>")
    h.append("</body></html>"); return "\n".join(h)

def html_to_pdf(pairs):
    js = """
const { chromium } = require('playwright');
(async () => {
  const b = await chromium.launch({ executablePath: '/opt/pw-browsers/chromium', args: ['--no-sandbox'] });
  for (const [src, out] of JSON.parse(process.argv[2])) {
    const p = await b.newPage(); await p.goto('file://' + src, { waitUntil: 'load' });
    await p.pdf({ path: out, format: 'A4', printBackground: true, margin: { top: '18mm', bottom: '18mm', left: '16mm', right: '16mm' },
      displayHeaderFooter: true, headerTemplate: '<span></span>',
      footerTemplate: '<div style="font-size:8px;color:#7B8A97;width:100%;text-align:center;font-family:sans-serif">Videl T&amp;L · Ex-Tarifário Factiun Sun BR · <span class="pageNumber"></span>/<span class="totalPages"></span></div>' });
    await p.close();
  }
  await b.close();
})().catch(e => { console.error(e); process.exit(1); });
"""
    tmp = ROOT / "_build" / "_pdf.js"; tmp.write_text(js)
    env = dict(os.environ, NODE_PATH=subprocess.check_output(["npm", "root", "-g"], text=True).strip())
    subprocess.run(["node", str(tmp), json.dumps(pairs)], check=True, env=env); tmp.unlink()

def main():
    pdf_all = "--pdf-all" in sys.argv
    pairs = []
    for md in sorted(ROOT.rglob("*.md")):
        if md.name == "README.md" and md.parent != ROOT: pass
        blocks = parse(md.read_text(encoding="utf-8"))
        title = next((b[2] for b in blocks if b[0] == "h"), md.stem)
        to_docx(blocks, md.with_suffix(".docx"))
        if pdf_all or md.name in PDF_DEFAULT:
            hp = ROOT / "_build" / (md.stem + ".html"); hp.write_text(to_html(blocks, title), encoding="utf-8")
            pairs.append([str(hp), str(md.with_suffix(".pdf"))])
        print("docx", md.relative_to(ROOT))
    if pairs:
        html_to_pdf(pairs)
        for src, out in pairs: pathlib.Path(src).unlink(); print("pdf ", pathlib.Path(out).relative_to(ROOT))

if __name__ == "__main__": main()
